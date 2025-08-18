#!/usr/bin/env python3
"""
PDF 및 DOCX 테스트 문서 생성
"""

import os
from pathlib import Path

# 테스트 디렉토리 생성
test_dir = Path("test_data")
test_dir.mkdir(exist_ok=True)

# DOCX 생성
try:
    from docx import Document
    
    print("Creating DOCX test file...")
    
    # DOCX 파일 생성
    docx_path = test_dir / "jewelry_appraisal.docx"
    doc = Document()
    
    # 제목
    doc.add_heading('주얼리 감정서 - Jewelry Appraisal Report', 0)
    
    # 기본 정보
    doc.add_heading('기본 정보 (Basic Information)', level=1)
    p = doc.add_paragraph()
    p.add_run('품목명 (Item): ').bold = True
    p.add_run('다이아몬드 목걸이 (Diamond Necklace)')
    
    p = doc.add_paragraph()
    p.add_run('브랜드 (Brand): ').bold = True
    p.add_run('루이까또즈 (Louis Cartier)')
    
    # 고객 상담 내용
    doc.add_heading('고객 상담 내용 (Customer Consultation)', level=1)
    
    consultation = """고객: 이 목걸이의 정확한 가치를 알고 싶어요. 할머니께서 물려주신 건데요.

감정사: 네, 꼼꼼히 감정해드렸습니다. 이 목걸이는 정말 훌륭한 작품이네요.
주 다이아몬드가 2캐럿이고 D컬러 VVS1 등급으로 최고급입니다.

고객: 보험을 위해서 감정서가 필요한데, 얼마나 걸리나요?

감정사: 공식 감정서는 3-5일 정도 소요됩니다. 보험용으로는 이 감정서로 충분하실 거예요.
현재 시장가치로는 약 1,250만원 정도로 평가됩니다."""
    
    doc.add_paragraph(consultation)
    
    doc.save(str(docx_path))
    print(f"DOCX created: {docx_path}")
    
except Exception as e:
    print(f"Error creating DOCX: {e}")

# PDF 생성 시도
try:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    
    print("Creating PDF test file...")
    
    # PDF 파일 생성
    pdf_path = test_dir / "jewelry_catalog.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    
    # 제목
    c.setFont("Helvetica-Bold", 16)
    c.drawString(100, 750, "Jewelry Catalog - Premium Collection")
    
    # 내용
    c.setFont("Helvetica", 12)
    content = [
        "Product: Diamond Ring Collection",
        "Category: Wedding Rings",
        "Description:",
        "- Material: 18K White Gold", 
        "- Main Stone: 1.5 Carat Diamond (VS1, F Color)",
        "- Price: $8,500",
        "Customer Review:",
        "Amazing ring! The diamond clarity is excellent."
    ]
    
    y_position = 700
    for line in content:
        c.drawString(100, y_position, line)
        y_position -= 25
    
    c.save()
    print(f"PDF created: {pdf_path}")
    
except ImportError:
    print("reportlab not available - creating simple PDF alternative")
    # 간단한 텍스트 기반 PDF 내용 생성
    pdf_content = """Jewelry Catalog - Premium Collection

Product: Diamond Ring Collection
Category: Wedding Rings

Description:
- Material: 18K White Gold
- Main Stone: 1.5 Carat Diamond (VS1, F Color)  
- Side Stones: 0.25 Carat total weight
- Setting: Prong Setting
- Price: $8,500

Customer Reviews:
고객: 정말 아름다운 반지입니다. 다이아몬드의 투명도가 뛰어나고 빛이 잘 반사되어 매우 만족합니다.

상담사: 이 제품은 저희의 베스트셀러 중 하나입니다. 특히 프로포즈용으로 많이 선택하시는 디자인이에요."""

    # PDF 대신 텍스트 파일로 생성
    pdf_alt_path = test_dir / "jewelry_catalog_pdf_content.txt"
    with open(pdf_alt_path, 'w', encoding='utf-8') as f:
        f.write(pdf_content)
    print(f"PDF content as text: {pdf_alt_path}")

print("Test document creation completed!")
