"""
솔로몬드 AI 시스템 - 이미지 처리 엔진
이미지, PDF, 문서에서 텍스트 추출 및 주얼리 특화 분석 모듈
"""

import os
import io
import tempfile
import base64
from pathlib import Path
from typing import Dict, List, Optional, Union, Tuple, Any
import asyncio
from concurrent.futures import ThreadPoolExecutor
import logging

# 이미지 처리 라이브러리
try:
    from PIL import Image, ImageEnhance, ImageFilter
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# OCR 라이브러리들
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
except ImportError:
    TESSERACT_AVAILABLE = False

try:
    import easyocr
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

# PDF 처리
try:
    import PyPDF2
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

# Word 문서 처리
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 주얼리 용어 데이터베이스 import
import json

class ImageProcessor:
    """이미지 및 문서 처리 클래스"""
    
    def __init__(self):
        self.supported_image_formats = ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif', '.webp']
        self.supported_document_formats = ['.pdf', '.docx', '.doc']
        self.ocr_languages = ['eng', 'kor', 'chi_sim', 'jpn']  # 영어, 한국어, 중국어(간체), 일본어
        
        # EasyOCR 초기화 (사용 가능한 경우)
        self.easyocr_reader = None
        if EASYOCR_AVAILABLE:
            try:
                self.easyocr_reader = easyocr.Reader(['en', 'ko', 'ch_sim', 'ja'])
            except Exception as e:
                logging.warning(f"EasyOCR 초기화 실패: {e}")
        
        # 주얼리 용어 데이터베이스 로드
        self.jewelry_terms = self._load_jewelry_terms()
        
        # 스레드 풀 executor
        self.executor = ThreadPoolExecutor(max_workers=4)
        
        logging.info(f"이미지 처리 엔진 초기화 완료")
        logging.info(f"지원 라이브러리: PIL={PIL_AVAILABLE}, Tesseract={TESSERACT_AVAILABLE}, EasyOCR={EASYOCR_AVAILABLE}")
        logging.info(f"문서 지원: PDF={PDF_AVAILABLE or PYMUPDF_AVAILABLE}, DOCX={DOCX_AVAILABLE}")
    
    def _load_jewelry_terms(self) -> Dict:
        """주얼리 용어 데이터베이스 로드"""
        try:
            # 상대 경로로 주얼리 용어 파일 찾기
            current_dir = Path(__file__).parent
            terms_file = current_dir.parent / "data" / "jewelry_terms.json"
            
            if terms_file.exists():
                with open(terms_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            else:
                logging.warning(f"주얼리 용어 파일을 찾을 수 없음: {terms_file}")
                return {}
        except Exception as e:
            logging.error(f"주얼리 용어 로드 실패: {e}")
            return {}
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """지원하는 파일 형식 반환"""
        return {
            "images": self.supported_image_formats,
            "documents": self.supported_document_formats,
            "all": self.supported_image_formats + self.supported_document_formats
        }
    
    def is_supported_file(self, filename: str) -> bool:
        """지원하는 파일인지 확인"""
        ext = Path(filename).suffix.lower()
        return ext in (self.supported_image_formats + self.supported_document_formats)
    
    def get_file_type(self, filename: str) -> str:
        """파일 타입 분류"""
        ext = Path(filename).suffix.lower()
        if ext in self.supported_image_formats:
            return "image"
        elif ext in self.supported_document_formats:
            return "document"
        else:
            return "unsupported"
    
    async def process_file(self, 
                          file_content: bytes, 
                          filename: str,
                          enhance_quality: bool = True,
                          ocr_method: str = "auto") -> Dict:
        """
        파일 처리 메인 함수
        
        Args:
            file_content: 파일 바이너리 데이터
            filename: 원본 파일명
            enhance_quality: 이미지 품질 향상 여부
            ocr_method: OCR 방법 ("tesseract", "easyocr", "auto")
            
        Returns:
            처리 결과 딕셔너리
        """
        try:
            file_type = self.get_file_type(filename)
            
            if file_type == "image":
                return await self.process_image(file_content, filename, enhance_quality, ocr_method)
            elif file_type == "document":
                return await self.process_document(file_content, filename)
            else:
                return {
                    "success": False,
                    "error": f"지원하지 않는 파일 형식: {Path(filename).suffix}",
                    "supported_formats": self.get_supported_formats()
                }
                
        except Exception as e:
            logging.error(f"파일 처리 오류: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename
            }
    
    async def process_image(self, 
                           image_content: bytes, 
                           filename: str,
                           enhance_quality: bool = True,
                           ocr_method: str = "auto") -> Dict:
        """
        이미지 처리 및 OCR
        
        Args:
            image_content: 이미지 바이너리 데이터
            filename: 원본 파일명
            enhance_quality: 이미지 품질 향상 여부
            ocr_method: OCR 방법
            
        Returns:
            처리 결과 딕셔너리
        """
        if not PIL_AVAILABLE:
            return {
                "success": False,
                "error": "PIL(Pillow) 라이브러리가 설치되지 않음. pip install Pillow로 설치하세요.",
                "filename": filename
            }
        
        try:
            print(f"🖼️ 이미지 처리 시작: {filename}")
            
            # PIL Image 객체 생성
            image = Image.open(io.BytesIO(image_content))
            
            # 이미지 정보 추출
            image_info = {
                "size": image.size,
                "mode": image.mode,
                "format": image.format,
                "file_size_mb": round(len(image_content) / (1024 * 1024), 2)
            }
            
            # 이미지 품질 향상 (선택적)
            if enhance_quality:
                image = await self._enhance_image_quality(image)
            
            # OCR 실행
            ocr_results = await self._perform_ocr(image, ocr_method)
            
            # 주얼리 특화 분석
            jewelry_analysis = self._analyze_jewelry_content(ocr_results.get("text", ""))
            
            # 결과 조합
            result = {
                "success": True,
                "filename": filename,
                "file_type": "image",
                "image_info": image_info,
                "ocr_results": ocr_results,
                "jewelry_analysis": jewelry_analysis,
                "processing_info": {
                    "enhanced": enhance_quality,
                    "ocr_method": ocr_results.get("method", "unknown")
                }
            }
            
            print(f"✅ 이미지 처리 완료: {len(ocr_results.get('text', ''))}자 추출")
            return result
            
        except Exception as e:
            logging.error(f"이미지 처리 오류: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename,
                "file_type": "image"
            }
    
    async def _enhance_image_quality(self, image: Image) -> Image:
        """이미지 품질 향상"""
        try:
            # RGB로 변환 (필요한 경우)
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # 해상도가 너무 낮으면 업스케일링
            width, height = image.size
            if width < 800 or height < 600:
                scale_factor = max(800/width, 600/height)
                new_size = (int(width * scale_factor), int(height * scale_factor))
                image = image.resize(new_size, Image.Resampling.LANCZOS)
            
            # 선명도 향상
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(1.2)
            
            # 대비 향상
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.1)
            
            # 노이즈 제거 (가벼운 필터)
            image = image.filter(ImageFilter.MedianFilter(size=3))
            
            return image
            
        except Exception as e:
            logging.warning(f"이미지 품질 향상 실패: {e}")
            return image
    
    async def _perform_ocr(self, image: Image, method: str = "auto") -> Dict:
        """OCR 실행"""
        results = {
            "text": "",
            "method": "none",
            "confidence": 0.0,
            "languages": [],
            "processing_time": 0.0
        }
        
        import time
        start_time = time.time()
        
        try:
            if method == "auto":
                # 사용 가능한 최선의 방법 자동 선택
                if EASYOCR_AVAILABLE and self.easyocr_reader:
                    method = "easyocr"
                elif TESSERACT_AVAILABLE:
                    method = "tesseract"
                else:
                    return {
                        **results,
                        "error": "사용 가능한 OCR 엔진이 없습니다. pytesseract 또는 easyocr을 설치하세요."
                    }
            
            if method == "easyocr" and EASYOCR_AVAILABLE and self.easyocr_reader:
                results = await self._ocr_with_easyocr(image)
            elif method == "tesseract" and TESSERACT_AVAILABLE:
                results = await self._ocr_with_tesseract(image)
            else:
                results["error"] = f"선택한 OCR 방법을 사용할 수 없습니다: {method}"
            
            results["processing_time"] = round(time.time() - start_time, 2)
            return results
            
        except Exception as e:
            logging.error(f"OCR 실행 오류: {e}")
            return {
                **results,
                "error": str(e),
                "processing_time": round(time.time() - start_time, 2)
            }
    
    async def _ocr_with_easyocr(self, image: Image) -> Dict:
        """EasyOCR을 사용한 텍스트 추출"""
        try:
            # PIL Image를 numpy array로 변환
            import numpy as np
            image_array = np.array(image)
            
            # 비동기로 OCR 실행
            loop = asyncio.get_event_loop()
            result = await loop.run_in_executor(
                self.executor, 
                self.easyocr_reader.readtext, 
                image_array
            )
            
            # 결과 파싱
            texts = []
            confidences = []
            
            for detection in result:
                if len(detection) >= 2:
                    text = detection[1]
                    confidence = detection[2] if len(detection) > 2 else 0.5
                    
                    texts.append(text)
                    confidences.append(confidence)
            
            full_text = " ".join(texts)
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
            
            return {
                "text": full_text,
                "method": "easyocr",
                "confidence": round(avg_confidence, 3),
                "languages": ["en", "ko", "ch_sim", "ja"],
                "details": result
            }
            
        except Exception as e:
            logging.error(f"EasyOCR 오류: {e}")
            return {
                "text": "",
                "method": "easyocr",
                "error": str(e)
            }
    
    async def _ocr_with_tesseract(self, image: Image) -> Dict:
        """Tesseract를 사용한 텍스트 추출"""
        try:
            # 임시 파일로 이미지 저장
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                image.save(temp_file.name, 'PNG')
                temp_path = temp_file.name
            
            try:
                # 비동기로 Tesseract 실행
                loop = asyncio.get_event_loop()
                
                # 다국어 설정
                lang_config = '+'.join(self.ocr_languages)
                
                text = await loop.run_in_executor(
                    self.executor,
                    lambda: pytesseract.image_to_string(
                        temp_path, 
                        lang=lang_config,
                        config='--oem 3 --psm 6'
                    )
                )
                
                # 신뢰도 정보 추출 (가능한 경우)
                try:
                    data = await loop.run_in_executor(
                        self.executor,
                        lambda: pytesseract.image_to_data(
                            temp_path,
                            lang=lang_config,
                            output_type=pytesseract.Output.DICT
                        )
                    )
                    
                    confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
                    avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
                    
                except Exception:
                    avg_confidence = 0.5  # 기본값
                
                return {
                    "text": text.strip(),
                    "method": "tesseract",
                    "confidence": round(avg_confidence / 100, 3),  # 0-1 범위로 정규화
                    "languages": self.ocr_languages
                }
                
            finally:
                # 임시 파일 정리
                try:
                    os.unlink(temp_path)
                except:
                    pass
                    
        except Exception as e:
            logging.error(f"Tesseract 오류: {e}")
            return {
                "text": "",
                "method": "tesseract", 
                "error": str(e)
            }
    
    async def process_document(self, 
                              doc_content: bytes, 
                              filename: str) -> Dict:
        """
        문서 처리 (PDF, Word 등)
        
        Args:
            doc_content: 문서 바이너리 데이터
            filename: 원본 파일명
            
        Returns:
            처리 결과 딕셔너리
        """
        try:
            ext = Path(filename).suffix.lower()
            
            if ext == '.pdf':
                return await self._process_pdf(doc_content, filename)
            elif ext in ['.docx', '.doc']:
                return await self._process_word(doc_content, filename)
            else:
                return {
                    "success": False,
                    "error": f"지원하지 않는 문서 형식: {ext}",
                    "filename": filename
                }
                
        except Exception as e:
            logging.error(f"문서 처리 오류: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename,
                "file_type": "document"
            }
    
    async def _process_pdf(self, pdf_content: bytes, filename: str) -> Dict:
        """PDF 파일 처리"""
        print(f"📄 PDF 처리 시작: {filename}")
        
        if not (PDF_AVAILABLE or PYMUPDF_AVAILABLE):
            return {
                "success": False,
                "error": "PDF 처리 라이브러리가 없습니다. PyPDF2 또는 PyMuPDF를 설치하세요.",
                "filename": filename
            }
        
        try:
            # PyMuPDF 우선 사용 (더 강력함)
            if PYMUPDF_AVAILABLE:
                return await self._process_pdf_with_pymupdf(pdf_content, filename)
            else:
                return await self._process_pdf_with_pypdf2(pdf_content, filename)
                
        except Exception as e:
            logging.error(f"PDF 처리 실패: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename,
                "file_type": "pdf"
            }
    
    async def _process_pdf_with_pymupdf(self, pdf_content: bytes, filename: str) -> Dict:
        """PyMuPDF를 사용한 PDF 처리"""
        try:
            doc = fitz.open(stream=pdf_content, filetype="pdf")
            
            pages_text = []
            total_text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                pages_text.append({
                    "page": page_num + 1,
                    "text": text,
                    "char_count": len(text)
                })
                total_text += f"\n=== 페이지 {page_num + 1} ===\n{text}\n"
            
            doc.close()
            
            # 주얼리 특화 분석
            jewelry_analysis = self._analyze_jewelry_content(total_text)
            
            result = {
                "success": True,
                "filename": filename,
                "file_type": "pdf",
                "text": total_text.strip(),
                "pages": pages_text,
                "page_count": len(pages_text),
                "total_chars": len(total_text),
                "jewelry_analysis": jewelry_analysis,
                "processing_info": {
                    "method": "pymupdf",
                    "ocr_used": False
                }
            }
            
            print(f"✅ PDF 처리 완료: {len(pages_text)}페이지, {len(total_text)}자")
            return result
            
        except Exception as e:
            logging.error(f"PyMuPDF 처리 오류: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename,
                "processing_method": "pymupdf"
            }
    
    async def _process_pdf_with_pypdf2(self, pdf_content: bytes, filename: str) -> Dict:
        """PyPDF2를 사용한 PDF 처리"""
        try:
            pdf_reader = PdfReader(io.BytesIO(pdf_content))
            
            pages_text = []
            total_text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                pages_text.append({
                    "page": page_num + 1,
                    "text": text,
                    "char_count": len(text)
                })
                total_text += f"\n=== 페이지 {page_num + 1} ===\n{text}\n"
            
            # 주얼리 특화 분석
            jewelry_analysis = self._analyze_jewelry_content(total_text)
            
            result = {
                "success": True,
                "filename": filename,
                "file_type": "pdf",
                "text": total_text.strip(),
                "pages": pages_text,
                "page_count": len(pages_text),
                "total_chars": len(total_text),
                "jewelry_analysis": jewelry_analysis,
                "processing_info": {
                    "method": "pypdf2",
                    "ocr_used": False
                }
            }
            
            print(f"✅ PDF 처리 완료: {len(pages_text)}페이지, {len(total_text)}자")
            return result
            
        except Exception as e:
            logging.error(f"PyPDF2 처리 오류: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename,
                "processing_method": "pypdf2"
            }
    
    async def _process_word(self, doc_content: bytes, filename: str) -> Dict:
        """Word 문서 처리"""
        print(f"📝 Word 문서 처리 시작: {filename}")
        
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "Word 문서 처리를 위해 python-docx가 필요합니다. pip install python-docx로 설치하세요.",
                "filename": filename
            }
        
        try:
            # 임시 파일로 저장 후 처리
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(doc_content)
                temp_path = temp_file.name
            
            try:
                # python-docx로 문서 읽기
                doc = Document(temp_path)
                
                paragraphs = []
                total_text = ""
                
                for i, paragraph in enumerate(doc.paragraphs):
                    text = paragraph.text.strip()
                    if text:  # 빈 문단 제외
                        paragraphs.append({
                            "paragraph": i + 1,
                            "text": text,
                            "char_count": len(text)
                        })
                        total_text += f"{text}\n"
                
                # 표 내용도 추출
                tables_text = []
                for table_num, table in enumerate(doc.tables):
                    table_content = []
                    for row in table.rows:
                        row_content = []
                        for cell in row.cells:
                            row_content.append(cell.text.strip())
                        table_content.append(row_content)
                    
                    if table_content:
                        tables_text.append({
                            "table": table_num + 1,
                            "content": table_content
                        })
                        
                        # 표 내용을 텍스트에 추가
                        for row in table_content:
                            total_text += " | ".join(row) + "\n"
                
                # 주얼리 특화 분석
                jewelry_analysis = self._analyze_jewelry_content(total_text)
                
                result = {
                    "success": True,
                    "filename": filename,
                    "file_type": "word",
                    "text": total_text.strip(),
                    "paragraphs": paragraphs,
                    "tables": tables_text,
                    "paragraph_count": len(paragraphs),
                    "table_count": len(tables_text),
                    "total_chars": len(total_text),
                    "jewelry_analysis": jewelry_analysis,
                    "processing_info": {
                        "method": "python-docx"
                    }
                }
                
                print(f"✅ Word 처리 완료: {len(paragraphs)}문단, {len(tables_text)}표, {len(total_text)}자")
                return result
                
            finally:
                # 임시 파일 정리
                try:
                    os.unlink(temp_path)
                except:
                    pass
                    
        except Exception as e:
            logging.error(f"Word 문서 처리 오류: {e}")
            return {
                "success": False,
                "error": str(e),
                "filename": filename,
                "file_type": "word"
            }
    
    def _analyze_jewelry_content(self, text: str) -> Dict:
        """주얼리 관련 내용 분석"""
        if not text:
            return {"terms_found": [], "categories": [], "insights": []}
        
        text_lower = text.lower()
        found_terms = []
        categories_found = set()
        
        # 주얼리 용어 데이터베이스에서 용어 검색
        for category, terms_data in self.jewelry_terms.items():
            if isinstance(terms_data, dict) and "terms" in terms_data:
                for term_info in terms_data["terms"]:
                    if isinstance(term_info, dict):
                        # 한국어, 영어, 중국어 용어 검색
                        for lang in ["korean", "english", "chinese"]:
                            if lang in term_info:
                                term = term_info[lang].lower()
                                if term in text_lower:
                                    found_terms.append({
                                        "term": term_info.get("korean", term),
                                        "category": category,
                                        "languages": {
                                            k: v for k, v in term_info.items() 
                                            if k in ["korean", "english", "chinese"]
                                        }
                                    })
                                    categories_found.add(category)
        
        # 주요 인사이트 생성
        insights = []
        
        if "diamonds" in categories_found or "보석류" in categories_found:
            insights.append("다이아몬드 관련 내용이 포함되어 있습니다.")
        
        if "pricing" in categories_found or "비즈니스" in categories_found:
            insights.append("가격 또는 비즈니스 관련 정보가 있습니다.")
        
        if "certification" in categories_found or "감정기관" in categories_found:
            insights.append("보석 감정 또는 인증 관련 내용입니다.")
        
        if len(found_terms) > 10:
            insights.append(f"풍부한 주얼리 전문 용어가 포함되어 있습니다. ({len(found_terms)}개 용어)")
        elif len(found_terms) > 5:
            insights.append("적당한 수준의 주얼리 전문 내용입니다.")
        elif len(found_terms) > 0:
            insights.append("일부 주얼리 관련 용어가 포함되어 있습니다.")
        else:
            insights.append("주얼리 전문 용어가 적게 포함되어 있습니다.")
        
        return {
            "terms_found": found_terms[:20],  # 최대 20개만 반환
            "total_terms": len(found_terms),
            "categories": list(categories_found),
            "insights": insights,
            "jewelry_relevance": "high" if len(found_terms) > 10 else "medium" if len(found_terms) > 3 else "low"
        }

# 전역 인스턴스
_image_processor_instance = None

def get_image_processor() -> ImageProcessor:
    """전역 이미지 처리기 인스턴스 반환"""
    global _image_processor_instance
    if _image_processor_instance is None:
        _image_processor_instance = ImageProcessor()
    return _image_processor_instance

# 편의 함수들
async def process_image_file(file_content: bytes, filename: str, **kwargs) -> Dict:
    """이미지 파일 처리 편의 함수"""
    processor = get_image_processor()
    return await processor.process_file(file_content, filename, **kwargs)

async def process_document_file(file_content: bytes, filename: str) -> Dict:
    """문서 파일 처리 편의 함수"""
    processor = get_image_processor()
    return await processor.process_file(file_content, filename)

def check_image_support() -> Dict:
    """이미지 처리 지원 상태 확인"""
    processor = get_image_processor()
    return {
        "supported_formats": processor.get_supported_formats(),
        "libraries": {
            "PIL": PIL_AVAILABLE,
            "tesseract": TESSERACT_AVAILABLE,
            "easyocr": EASYOCR_AVAILABLE,
            "pypdf2": PDF_AVAILABLE,
            "pymupdf": PYMUPDF_AVAILABLE,
            "docx": DOCX_AVAILABLE
        },
        "ocr_methods": ["tesseract", "easyocr", "auto"],
        "languages": ["eng", "kor", "chi_sim", "jpn"]
    }

if __name__ == "__main__":
    # 테스트 코드
    print("이미지 처리 엔진 테스트")
    support_info = check_image_support()
    print(f"지원 상태: {support_info}")
